"""Business logic for the Leads admin group (B.7.5d — Export CSV + DOCX).

Port of:
- app/api/admin/leads-chatwit/export-csv/route.ts (GET)
- app/api/admin/leads-chatwit/export-docx/route.ts (POST)
"""

from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from domains.socialwise.db.models.arquivo_lead_oab import ArquivoLeadOab
from domains.socialwise.db.models.lead import Lead
from domains.socialwise.db.models.lead_oab_data import LeadOabData
from domains.socialwise.db.models.user import User
from domains.socialwise.db.models.usuario_chatwit import UsuarioChatwit
from platform_core.logging.config import get_logger

logger = get_logger(__name__)


class ExportServiceError(Exception):
    pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _escape_csv(value: Any) -> str:
    """Escape a value for CSV output."""
    if value is None:
        return ""
    s = str(value)
    if "," in s or '"' in s or "\n" in s:
        return f'"{s.replace(chr(34), chr(34) + chr(34))}"'
    return s


def _format_date(dt: datetime | None) -> str:
    """Format a datetime to pt-BR locale string."""
    if not dt:
        return ""
    return dt.strftime("%d/%m/%Y %H:%M:%S")


def _count_files_by_type(arquivos: list[ArquivoLeadOab]) -> dict[str, int]:
    """Count files by type: total, pdf, image."""
    total = len(arquivos)
    pdf = sum(1 for a in arquivos if a.file_type and "pdf" in a.file_type.lower())
    image = sum(
        1
        for a in arquivos
        if a.file_type
        and any(t in a.file_type.lower() for t in ("image", "jpg", "png"))
    )
    return {"total": total, "pdf": pdf, "image": image}


def _bool_pt(value: bool | None) -> str:
    return "Sim" if value else "Não"


# ---------------------------------------------------------------------------
# 1. Export CSV
# ---------------------------------------------------------------------------


async def export_csv(
    session: AsyncSession,
    user_id: str,
    user_role: str,
    search_term: str | None = None,
) -> str:
    """Export leads to CSV format.

    Returns the CSV content as a string (with UTF-8 BOM).
    """
    # Build where conditions
    conditions = []

    # Role-based access control
    if user_role != "SUPERADMIN":
        # Check if user has a Chatwit token
        usr_result = await session.execute(
            select(UsuarioChatwit.chatwit_access_token)
            .where(UsuarioChatwit.app_user_id == user_id)
        )
        token = usr_result.scalar_one_or_none()
        if not token:
            raise ExportServiceError("Nenhum lead disponível para exportação")

        conditions.append(UsuarioChatwit.app_user_id == user_id)

    # Search filter
    search_filters = []
    if search_term:
        cleaned = "".join(c for c in search_term if c.isalnum() or c in ("@", ".", "-"))
        like_term = f"%{search_term}%"
        like_cleaned = f"%{cleaned}%"

        search_filters = [
            Lead.name.ilike(like_term),
            LeadOabData.nome_real.ilike(like_term),
            Lead.phone.ilike(like_term),
            Lead.email.ilike(like_term),
            LeadOabData.id.ilike(like_term),
            LeadOabData.lead_id.ilike(like_term),
        ]
        if cleaned != search_term:
            search_filters.append(Lead.phone.ilike(like_cleaned))

    # Build query
    from sqlalchemy import or_

    query = (
        select(LeadOabData)
        .join(Lead, LeadOabData.lead_id == Lead.id)
        .outerjoin(UsuarioChatwit, LeadOabData.usuario_chatwit_id == UsuarioChatwit.id)
        .options(
            selectinload(LeadOabData.lead),
            selectinload(LeadOabData.usuario_chatwit),
            selectinload(LeadOabData.arquivos),
        )
        .order_by(Lead.updated_at.desc())
    )

    if conditions:
        query = query.where(*conditions)
    if search_filters:
        query = query.where(or_(*search_filters))

    result = await session.execute(query)
    leads = result.scalars().all()

    # CSV headers
    headers = [
        "ID", "Nome", "Nome Real", "Telefone", "Email",
        "Usuário", "Canal", "Status", "Concluído", "Fez Recurso",
        "Datas Recurso", "Total Arquivos", "Arquivos PDF", "Arquivos Imagem",
        "PDF Unificado", "Imagens Convertidas", "Prova Manuscrita",
        "Manuscrito Processado", "Aguardando Manuscrito", "Espelho Processado",
        "Aguardando Espelho", "Análise Processada", "Aguardando Análise",
        "Análise Validada", "Consultoria Fase 2", "Especialidade",
        "Seccional", "Área Jurídica", "Nota Final", "Situação",
        "Inscrição", "Exames Participados", "Observações",
        "Data Criação", "Última Atualização",
    ]

    rows: list[str] = []
    for ld in leads:
        lead_data = ld.lead
        nome_real = ld.nome_real
        if not nome_real or nome_real == "undefined":
            nome_real = (lead_data.name if lead_data else None) or "Nome não informado"

        file_stats = _count_files_by_type(ld.arquivos or [])

        # Parse datasRecurso
        datas_recurso = ""
        if ld.datas_recurso:
            try:
                datas = json.loads(ld.datas_recurso)
                datas_recurso = "; ".join(datas) if isinstance(datas, list) else ""
            except (json.JSONDecodeError, TypeError):
                pass

        # Count imagens convertidas
        imgs_count = 0
        if ld.imagens_convertidas:
            try:
                imgs = json.loads(ld.imagens_convertidas)
                imgs_count = len(imgs) if isinstance(imgs, list) else 0
            except (json.JSONDecodeError, TypeError):
                pass

        # Parse examesParticipados
        exames = ""
        if ld.exames_participados:
            if isinstance(ld.exames_participados, (list, dict)):
                exames = json.dumps(ld.exames_participados, ensure_ascii=False)
            else:
                exames = str(ld.exames_participados)

        usr = ld.usuario_chatwit
        row_values = [
            _escape_csv(ld.id),
            _escape_csv(lead_data.name if lead_data else ""),
            _escape_csv(nome_real),
            _escape_csv(lead_data.phone if lead_data else ""),
            _escape_csv(lead_data.email if lead_data else ""),
            _escape_csv(usr.name if usr else ""),
            _escape_csv(usr.channel if usr else ""),
            _escape_csv(ld.situacao or "Pendente"),
            _escape_csv(_bool_pt(ld.concluido)),
            _escape_csv(_bool_pt(ld.fez_recurso)),
            _escape_csv(datas_recurso),
            _escape_csv(file_stats["total"]),
            _escape_csv(file_stats["pdf"]),
            _escape_csv(file_stats["image"]),
            _escape_csv(_bool_pt(bool(ld.pdf_unificado))),
            _escape_csv(str(imgs_count) if imgs_count > 0 else "Não"),
            _escape_csv(_bool_pt(bool(ld.prova_manuscrita))),
            _escape_csv(_bool_pt(ld.manuscrito_processado)),
            _escape_csv(_bool_pt(ld.aguardando_manuscrito)),
            _escape_csv(_bool_pt(ld.espelho_processado)),
            _escape_csv(_bool_pt(ld.aguardando_espelho)),
            _escape_csv(_bool_pt(ld.analise_processada)),
            _escape_csv(_bool_pt(ld.aguardando_analise)),
            _escape_csv(_bool_pt(ld.analise_validada)),
            _escape_csv(_bool_pt(ld.consultoria_fase2)),
            _escape_csv(ld.especialidade or ""),
            _escape_csv(ld.seccional or ""),
            _escape_csv(ld.area_juridica or ""),
            _escape_csv(ld.nota_final or ""),
            _escape_csv(ld.situacao or ""),
            _escape_csv(ld.inscricao or ""),
            _escape_csv(exames),
            _escape_csv(ld.anotacoes or ""),
            _escape_csv(_format_date(lead_data.created_at if lead_data else None)),
            _escape_csv(_format_date(lead_data.updated_at if lead_data else None)),
        ]
        rows.append(",".join(row_values))

    csv_content = ",".join(headers) + "\n" + "\n".join(rows)
    bom = "\ufeff"

    logger.info("csv_export", leads_count=len(leads), user_id=user_id)
    return bom + csv_content


# ---------------------------------------------------------------------------
# 2. Export DOCX
# ---------------------------------------------------------------------------


def export_html_to_docx(html: str, lead_id: str) -> bytes:
    """Convert HTML content to DOCX format.

    Port of: app/api/admin/leads-chatwit/export-docx/route.ts
    Uses python-docx to create a simple document from HTML.
    """
    from html.parser import HTMLParser

    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    class _HtmlToDocxParser(HTMLParser):
        """Simple HTML parser that extracts text into docx paragraphs."""

        def __init__(self, document: DocxDocument):
            super().__init__()
            self._doc = document
            self._current_text = ""
            self._heading_level = 0
            self._in_list = False
            self._list_counter = 0
            self._bold = False
            self._italic = False

        def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
            tag = tag.lower()
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._flush()
                self._heading_level = int(tag[1])
            elif tag == "br":
                self._current_text += "\n"
            elif tag == "p":
                self._flush()
            elif tag == "li":
                self._flush()
                self._in_list = True
                self._list_counter += 1
            elif tag in ("strong", "b"):
                self._bold = True
            elif tag in ("em", "i"):
                self._italic = True
            elif tag in ("ul", "ol"):
                self._list_counter = 0

        def handle_endtag(self, tag: str):
            tag = tag.lower()
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                text = self._current_text.strip()
                if text:
                    self._doc.add_heading(text, level=self._heading_level)
                self._current_text = ""
                self._heading_level = 0
            elif tag == "p":
                self._flush()
            elif tag == "li":
                text = self._current_text.strip()
                if text:
                    self._doc.add_paragraph(text, style="List Bullet")
                self._current_text = ""
                self._in_list = False
            elif tag in ("strong", "b"):
                self._bold = False
            elif tag in ("em", "i"):
                self._italic = False

        def handle_data(self, data: str):
            self._current_text += data

        def _flush(self):
            text = self._current_text.strip()
            if text:
                para = self._doc.add_paragraph()
                run = para.add_run(text)
                run.bold = self._bold
                run.italic = self._italic
            self._current_text = ""

        def close(self):
            self._flush()
            super().close()

    parser = _HtmlToDocxParser(doc)
    parser.feed(html)
    parser.close()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
