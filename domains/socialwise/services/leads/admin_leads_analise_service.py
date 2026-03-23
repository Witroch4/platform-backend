"""Business logic for the Leads admin group (B.7.5c — Análise + Recurso).

Port of:
- app/api/admin/leads-chatwit/enviar-analise/route.ts (POST)
- app/api/admin/leads-chatwit/enviar-analise-validada/route.ts (POST)
- app/api/admin/leads-chatwit/gerar-recurso-interno/route.ts (POST)
- app/api/admin/leads-chatwit/enviar-recurso-validado/route.ts (POST)
- app/api/admin/leads-chatwit/enviar-consultoriafase2/route.ts (POST)
"""

from __future__ import annotations

import asyncio
import io
import json
import os
import re
from html.parser import HTMLParser
from typing import Any
from urllib.parse import urlparse

import httpx
from sqlalchemy import select, or_, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from domains.socialwise.db.models.arquivo_lead_oab import ArquivoLeadOab
from domains.socialwise.db.models.espelho_biblioteca import EspelhoBiblioteca
from domains.socialwise.db.models.lead import Lead
from domains.socialwise.db.models.lead_oab_data import LeadOabData
from domains.socialwise.db.models.usuario_chatwit import UsuarioChatwit
from platform_core.config import settings
from platform_core.logging.config import get_logger

logger = get_logger(__name__)


class AnaliseServiceError(Exception):
    pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _to_str(value: Any) -> str:
    """Convert any value to string for LLM input."""
    if isinstance(value, str):
        return value
    if value is None:
        return ""
    return json.dumps(value, ensure_ascii=False)


def _optimize_mirror_payload(espelho: Any) -> Any:
    """Optimize mirror payload to reduce token size for external webhook.

    Port of: lib/oab-eval/mirror-formatter.ts → optimizeMirrorPayload()
    Removes verbose fields not needed for analysis (descriptions, original images, etc.)
    """
    if espelho is None:
        return None

    if isinstance(espelho, str):
        try:
            espelho = json.loads(espelho)
        except json.JSONDecodeError:
            return espelho

    if not isinstance(espelho, dict):
        return espelho

    # Validate structure
    if not all(k in espelho for k in ("meta", "aluno", "itens")):
        return espelho

    # Build optimized copy keeping only essential fields
    optimized: dict[str, Any] = {}

    if "meta" in espelho:
        optimized["meta"] = espelho["meta"]
    if "aluno" in espelho:
        optimized["aluno"] = espelho["aluno"]
    if "totais" in espelho:
        optimized["totais"] = espelho["totais"]

    # For itens, keep score data but strip verbose descriptions
    if "itens" in espelho and isinstance(espelho["itens"], list):
        optimized_itens = []
        for item in espelho["itens"]:
            if isinstance(item, dict):
                opt_item = {
                    k: v
                    for k, v in item.items()
                    if k not in ("descricao_longa", "imagem_original", "raw_html")
                }
                optimized_itens.append(opt_item)
            else:
                optimized_itens.append(item)
        optimized["itens"] = optimized_itens

    return optimized


def _parse_espelho_images(espelho_correcao: str | None) -> list[str]:
    """Parse espelhoCorrecao JSON string to list of image URLs."""
    if not espelho_correcao:
        return []
    try:
        parsed = json.loads(espelho_correcao)
        return parsed if isinstance(parsed, list) else []
    except (json.JSONDecodeError, TypeError):
        return []


def _format_text_value(value: Any, prefix: str = "") -> str:
    """Format a prova/espelho value handling string, array, and object types.

    Port of the multi-format handling in enviar-consultoriafase2/route.ts.
    """
    if value is None:
        return ""
    if isinstance(value, str):
        return f"{prefix}\n{value}" if prefix else value
    if isinstance(value, list):
        parts = []
        for item in value:
            if isinstance(item, dict) and "output" in item:
                parts.append(item["output"])
            else:
                parts.append(json.dumps(item, ensure_ascii=False))
        joined = "\n\n---------------------------------\n\n".join(parts)
        return f"{prefix}\n{joined}" if prefix else joined
    if isinstance(value, dict):
        text = json.dumps(value, ensure_ascii=False, indent=2)
        return f"{prefix}\n{text}" if prefix else text
    return str(value)


def _extract_conversation_id(lead_url: str) -> str:
    """Extract conversationId from Chatwit URL.

    e.g. https://chatwit.witdev.com.br/app/accounts/3/conversations/1199 → '1199'
    """
    parsed = urlparse(lead_url)
    parts = parsed.path.strip("/").split("/")
    try:
        conv_idx = parts.index("conversations")
        return parts[conv_idx + 1]
    except (ValueError, IndexError):
        raise AnaliseServiceError(f"leadUrl fora do formato esperado: {lead_url}")


# ---------------------------------------------------------------------------
# HTML → DOCX converter (lightweight, no extra deps beyond python-docx)
# ---------------------------------------------------------------------------


class _SimpleHTMLToDocx(HTMLParser):
    """Minimal HTML parser that creates python-docx paragraphs."""

    def __init__(self, document: Any):
        super().__init__()
        self.doc = document
        self._current_paragraph = None
        self._bold = False
        self._underline = False
        self._heading_level = 0
        self._text_buf: list[str] = []

    def _flush_text(self) -> None:
        text = "".join(self._text_buf).strip()
        if not text:
            self._text_buf.clear()
            return
        if self._current_paragraph is None:
            if self._heading_level:
                self._current_paragraph = self.doc.add_heading(text, level=self._heading_level)
            else:
                self._current_paragraph = self.doc.add_paragraph()
                run = self._current_paragraph.add_run(text)
                run.bold = self._bold
                run.underline = self._underline
        else:
            run = self._current_paragraph.add_run(text)
            run.bold = self._bold
            run.underline = self._underline
        self._text_buf.clear()

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in ("h1", "h2", "h3"):
            self._flush_text()
            self._heading_level = int(tag[1])
            self._current_paragraph = None
        elif tag == "b" or tag == "strong":
            self._flush_text()
            self._bold = True
        elif tag == "u":
            self._flush_text()
            self._underline = True
        elif tag == "p":
            self._flush_text()
            self._current_paragraph = None
        elif tag == "br":
            self._flush_text()
            if self._current_paragraph:
                self._current_paragraph.add_run("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in ("h1", "h2", "h3"):
            self._flush_text()
            self._heading_level = 0
            self._current_paragraph = None
        elif tag == "b" or tag == "strong":
            self._flush_text()
            self._bold = False
        elif tag == "u":
            self._flush_text()
            self._underline = False
        elif tag == "p":
            self._flush_text()
            self._current_paragraph = None

    def handle_data(self, data: str) -> None:
        self._text_buf.append(data)

    def close(self) -> None:
        self._flush_text()
        super().close()


def _html_to_docx_bytes(html: str) -> bytes:
    """Convert HTML string to DOCX bytes using python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    parser = _SimpleHTMLToDocx(doc)
    parser.feed(html)
    parser.close()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 1. enviar-analise (POST)
# ---------------------------------------------------------------------------


async def enviar_analise(session: AsyncSession, payload: dict[str, Any]) -> dict[str, Any]:
    """Send lead for analysis — internal (TaskIQ) or external (webhook).

    Port of: app/api/admin/leads-chatwit/enviar-analise/route.ts
    """
    lead_id = payload.get("leadId") or payload.get("leadID")
    if not lead_id:
        raise AnaliseServiceError("leadId é obrigatório")

    selected_provider = payload.get("selectedProvider")
    source_id = payload.get("sourceId")

    # Fetch lead
    stmt = (
        select(LeadOabData)
        .options(
            selectinload(LeadOabData.lead),
        )
        .where(LeadOabData.id == lead_id)
    )
    result = await session.execute(stmt)
    lead_oab = result.scalar_one_or_none()

    if not lead_oab:
        raise AnaliseServiceError("Lead não encontrado")

    # Fetch arquivos
    arquivos_stmt = select(ArquivoLeadOab).where(ArquivoLeadOab.lead_oab_data_id == lead_id)
    arquivos_result = await session.execute(arquivos_stmt)
    arquivos = list(arquivos_result.scalars().all())

    lead_name = lead_oab.nome_real or (lead_oab.lead.name if lead_oab.lead else None) or "Lead sem nome"
    lead_phone = lead_oab.lead.phone if lead_oab.lead else None

    # Optimize espelho
    espelho_otimizado = lead_oab.texto_do_espelho
    if espelho_otimizado:
        try:
            espelho_parsed = espelho_otimizado if isinstance(espelho_otimizado, dict) else json.loads(str(espelho_otimizado))
            if isinstance(espelho_parsed, dict) and all(k in espelho_parsed for k in ("meta", "aluno", "itens")):
                espelho_otimizado = _optimize_mirror_payload(espelho_parsed)
                logger.info("espelho_optimized", lead_id=lead_id)
            else:
                logger.warning("espelho_structure_invalid", lead_id=lead_id)
        except (json.JSONDecodeError, TypeError):
            logger.info("espelho_not_json", lead_id=lead_id)

    # Feature flag: internal analysis
    use_internal = settings.oab_agent_local

    if use_internal:
        logger.info("enviar_analise_internal", lead_id=lead_id)

        texto_prova = _to_str(lead_oab.prova_manuscrita)
        texto_espelho = _to_str(espelho_otimizado)

        if not texto_prova or len(texto_prova.strip()) < 10:
            raise AnaliseServiceError("Texto da prova manuscrita ausente ou muito curto.")

        if not texto_espelho or len(texto_espelho.strip()) < 10:
            raise AnaliseServiceError("Texto do espelho ausente ou muito curto.")

        # Enqueue TaskIQ job
        from domains.socialwise.tasks.analysis_generation import process_analysis_generation_task

        task_payload = {
            "leadId": lead_id,
            "textoProva": texto_prova,
            "textoEspelho": texto_espelho,
            "selectedProvider": selected_provider,
            "telefone": lead_phone,
            "nome": lead_name,
        }

        task = await process_analysis_generation_task.kiq(task_payload)

        # Mark as waiting
        await session.execute(
            update(LeadOabData)
            .where(LeadOabData.id == lead_id)
            .values(aguardando_analise=True)
        )
        await session.commit()

        return {
            "success": True,
            "message": "Análise enviada para processamento interno",
            "jobId": str(task.task_id) if hasattr(task, "task_id") else "queued",
            "internal": True,
            "operation": {
                "leadId": lead_id,
                "stage": "analysis",
            },
            "_status_code": 202,
        }

    # External webhook path
    logger.info("enviar_analise_external", lead_id=lead_id)

    webhook_url = settings.socialwise_webhook_url
    if not webhook_url:
        raise AnaliseServiceError("URL do webhook não configurada no ambiente")

    espelho_str = _to_str(espelho_otimizado)

    webhook_payload: dict[str, Any] = {
        "leadID": lead_id,
        "nome": lead_name,
        "telefone": lead_phone,
        "analise": True,
        "arquivos": [
            {"id": a.id, "url": a.data_url, "tipo": a.file_type, "nome": a.file_type}
            for a in arquivos
        ],
        "arquivos_pdf": (
            [{"id": lead_id, "url": lead_oab.pdf_unificado, "nome": "PDF Unificado"}]
            if lead_oab.pdf_unificado
            else []
        ),
        "textoManuscrito": _to_str(lead_oab.prova_manuscrita),
        "textoEspelho": espelho_str,
        "metadata": {
            "leadUrl": lead_oab.lead_url,
            "sourceId": (lead_oab.lead.source_identifier if lead_oab.lead else None) or source_id,
            "concluido": lead_oab.concluido,
            "fezRecurso": lead_oab.fez_recurso,
        },
    }

    # Add espelho images if exist
    espelho_images = _parse_espelho_images(lead_oab.espelho_correcao)
    if espelho_images:
        webhook_payload["arquivos_imagens_espelho"] = [
            {"id": f"{lead_id}-espelho-{i}", "url": url, "nome": f"Espelho {i + 1}"}
            for i, url in enumerate(espelho_images)
        ]

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.post(webhook_url, json=webhook_payload)
        if resp.status_code >= 400:
            raise AnaliseServiceError(f"Erro ao enviar análise: {resp.status_code}")

    # Mark as waiting
    await session.execute(
        update(LeadOabData)
        .where(LeadOabData.id == lead_id)
        .values(aguardando_analise=True)
    )
    await session.commit()

    return {"success": True, "message": "Lead enviado para análise com sucesso"}


# ---------------------------------------------------------------------------
# 2. enviar-analise-validada (POST)
# ---------------------------------------------------------------------------


async def enviar_analise_validada(session: AsyncSession, payload: dict[str, Any]) -> dict[str, Any]:
    """Send validated analysis for PDF generation.

    Port of: app/api/admin/leads-chatwit/enviar-analise-validada/route.ts
    """
    lead_id = payload.get("leadID")
    if not lead_id:
        raise AnaliseServiceError("leadID não fornecido")

    # Fetch lead
    stmt = (
        select(LeadOabData)
        .options(selectinload(LeadOabData.lead))
        .where(LeadOabData.id == lead_id)
    )
    result = await session.execute(stmt)
    lead_oab = result.scalar_one_or_none()

    if not lead_oab:
        raise AnaliseServiceError("Lead não encontrado")

    # Mark as validated + save analysis data
    analise_data = payload.get("analiseData", {})
    await session.execute(
        update(LeadOabData)
        .where(LeadOabData.id == lead_id)
        .values(
            analise_validada=True,
            analise_preliminar=analise_data,
        )
    )
    await session.commit()

    # Detect simulado
    is_simulado = analise_data.get("analisesimuladovalidado") is True

    # Build request payload
    lead_name = lead_oab.nome_real or (lead_oab.lead.name if lead_oab.lead else None) or ""
    lead_phone = lead_oab.lead.phone if lead_oab.lead else None

    request_payload: dict[str, Any] = {
        "leadID": lead_id,
        "telefone": lead_phone,
    }

    # Flag type
    if is_simulado:
        request_payload["analisesimuladovalidado"] = True
    else:
        request_payload["analisevalidada"] = True

    # Header fields
    for field in (
        "exameDescricao", "inscricao", "seccional", "areaJuridica",
        "notaFinal", "situacao", "subtotalPeca", "subtotalQuestoes", "conclusao",
    ):
        request_payload[field] = analise_data.get(field, "")

    request_payload["nomeExaminando"] = analise_data.get("nomeExaminando") or lead_name
    request_payload["pontosPeca"] = analise_data.get("pontosPeca", [])
    request_payload["pontosQuestoes"] = analise_data.get("pontosQuestoes", [])
    request_payload["argumentacao"] = analise_data.get("argumentacao", [])

    # Merge remaining analysis data (excluding control flags)
    for k, v in analise_data.items():
        if k not in ("analisesimuladovalidado", "analiseValidada") and k not in request_payload:
            request_payload[k] = v

    use_internal = settings.analise_validada_interna

    if use_internal:
        logger.info("analise_validada_internal", lead_id=lead_id)

        from domains.socialwise.tasks.lead_cells import process_lead_cell_task

        task_payload = {
            "type": "analise",
            "leadID": lead_id,
            "analiseData": request_payload,
            "generatePdfInternally": True,
            "analiseValidada": None if is_simulado else True,
            "analiseSimuladoValidada": True if is_simulado else None,
            "nome": request_payload.get("nomeExaminando"),
            "telefone": lead_phone,
        }

        await process_lead_cell_task.kiq(task_payload)
        logger.info("analise_validada_job_queued", lead_id=lead_id)
    else:
        # Fire-and-forget to external webhook
        webhook_url = settings.socialwise_webhook_url
        if not webhook_url:
            raise AnaliseServiceError("URL do webhook não configurada")

        logger.info("analise_validada_external", lead_id=lead_id, webhook_url=webhook_url)

        # Fire-and-forget (don't await result)
        async def _send_webhook() -> None:
            try:
                async with httpx.AsyncClient(timeout=30) as client:
                    resp = await client.post(webhook_url, json=request_payload)
                    if resp.status_code >= 400:
                        logger.error("analise_validada_webhook_error", status=resp.status_code)
                    else:
                        logger.info("analise_validada_webhook_sent", lead_id=lead_id)
            except Exception as exc:
                logger.error("analise_validada_webhook_failed", error=str(exc))

        asyncio.create_task(_send_webhook())

    return {"success": True, "message": "Análise validada enviada com sucesso"}


# ---------------------------------------------------------------------------
# 3. gerar-recurso-interno (POST)
# ---------------------------------------------------------------------------


async def gerar_recurso_interno(session: AsyncSession, payload: dict[str, Any]) -> dict[str, Any]:
    """Generate recurso via AI agent.

    Port of: app/api/admin/leads-chatwit/gerar-recurso-interno/route.ts
    """
    lead_id = payload.get("leadId")
    analise_validada = payload.get("analiseValidada")
    dados_adicionais = payload.get("dadosAdicionais")
    selected_provider = payload.get("selectedProvider")

    if not lead_id:
        raise AnaliseServiceError("O ID do lead é obrigatório.")

    if not analise_validada:
        raise AnaliseServiceError("A análise validada é obrigatória para gerar o recurso.")

    logger.info("gerar_recurso_start", lead_id=lead_id)

    from domains.socialwise.services.oab_eval.recurso_agent import run_recurso

    result = await run_recurso(
        session,
        lead_id=lead_id,
        analise_validada=analise_validada,
        dados_adicionais=dados_adicionais,
        selected_provider=selected_provider,
    )

    if not result.get("success"):
        return {
            "success": False,
            "error": result.get("error", "Falha ao gerar o recurso."),
            "_status_code": 500,
        }

    # Save result to DB — try both id and leadId
    stmt = (
        select(LeadOabData)
        .where(or_(LeadOabData.id == lead_id, LeadOabData.lead_id == lead_id))
    )
    db_result = await session.execute(stmt)
    lead_oab = db_result.scalar_one_or_none()

    if lead_oab:
        await session.execute(
            update(LeadOabData)
            .where(LeadOabData.id == lead_oab.id)
            .values(
                recurso_preliminar=result.get("recursoOutput") or {},
                aguardando_recurso=False,
            )
        )
        await session.commit()
        logger.info("recurso_saved", lead_oab_id=lead_oab.id)
    else:
        logger.warning("recurso_lead_not_found", lead_id=lead_id)

    return {
        "success": True,
        "recursoOutput": result.get("recursoOutput"),
        "message": "Recurso gerado com sucesso.",
    }


# ---------------------------------------------------------------------------
# 4. enviar-recurso-validado (POST)
# ---------------------------------------------------------------------------


async def enviar_recurso_validado(session: AsyncSession, payload: dict[str, Any]) -> dict[str, Any]:
    """Validate recurso, generate DOCX, send to Chatwit chat.

    Port of: app/api/admin/leads-chatwit/enviar-recurso-validado/route.ts
    """
    lead_id = payload.get("leadID")
    html_content = payload.get("html")

    if not lead_id:
        raise AnaliseServiceError("leadID não fornecido")
    if not html_content:
        raise AnaliseServiceError("html do recurso não fornecido")

    # 1) Fetch lead
    stmt = (
        select(LeadOabData)
        .options(selectinload(LeadOabData.lead))
        .where(LeadOabData.id == lead_id)
    )
    result = await session.execute(stmt)
    lead_oab = result.scalar_one_or_none()

    if not lead_oab:
        raise AnaliseServiceError("Lead não encontrado")
    if not lead_oab.lead_url:
        raise AnaliseServiceError("Lead sem leadUrl (conversa não vinculada)")

    # 2) Find accountId from UsuarioChatwit
    ucw_stmt = (
        select(UsuarioChatwit)
        .join(
            LeadOabData,
            LeadOabData.usuario_chatwit_id == UsuarioChatwit.id,
        )
        .where(LeadOabData.id == lead_id)
    )
    ucw_result = await session.execute(ucw_stmt)
    usuario_chatwit = ucw_result.scalar_one_or_none()

    if not usuario_chatwit or not usuario_chatwit.chatwit_account_id:
        raise AnaliseServiceError("Usuário Chatwit não configurado")

    # 3) Access token
    access_token = (
        payload.get("accessToken")
        or settings.chatwit_access_token
    )
    if not access_token:
        raise AnaliseServiceError("Token de acesso não configurado")

    # 4) Generate DOCX from HTML
    docx_bytes = _html_to_docx_bytes(html_content)
    logger.info("docx_generated", lead_id=lead_id, size=len(docx_bytes))

    # 5) Send DOCX to Chatwit API
    conversation_id = _extract_conversation_id(lead_oab.lead_url)
    account_id = usuario_chatwit.chatwit_account_id
    message = payload.get("message") or "Segue o nosso Recurso, qualquer dúvida estamos à disposição."
    nome_real = lead_oab.nome_real or (lead_oab.lead.name if lead_oab.lead else None) or "lead"
    filename = f"recurso_{re.sub(r'[^a-zA-Z0-9]', '_', nome_real)}.docx"

    chatwit_base = settings.chatwit_base_url or "https://chatwit.witdev.com.br"
    chatwit_url = f"{chatwit_base}/api/v1/accounts/{account_id}/conversations/{conversation_id}/messages"

    async with httpx.AsyncClient(timeout=60) as client:
        files = {
            "attachments[]": (
                filename,
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        }
        data = {
            "content": message,
            "message_type": "outgoing",
        }

        resp = await client.post(
            chatwit_url,
            data=data,
            files=files,
            headers={"api_access_token": access_token},
        )

        if resp.status_code >= 400:
            logger.error("chatwit_send_failed", status=resp.status_code, body=resp.text[:200])
            raise AnaliseServiceError(f"Falha ao enviar DOCX para Chatwit: {resp.status_code}")

        chatwit_response = resp.json()

    logger.info("docx_sent_to_chatwit", conversation_id=conversation_id)

    # 6) Update DB
    update_data: dict[str, Any] = {
        "recurso_validado": True,
        "fez_recurso": True,
        "aguardando_recurso": False,
        "anotacoes": message,
    }
    if payload.get("textoRecurso"):
        update_data["recurso_preliminar"] = {"textoRecurso": payload["textoRecurso"]}

    await session.execute(
        update(LeadOabData)
        .where(LeadOabData.id == lead_id)
        .values(**update_data)
    )
    await session.commit()

    logger.info("recurso_validado_db_updated", lead_id=lead_id)

    return {
        "success": True,
        "message": "Recurso validado e enviado para o chat com sucesso",
        "chatwoot": chatwit_response,
    }


# ---------------------------------------------------------------------------
# 5. enviar-consultoriafase2 (POST)
# ---------------------------------------------------------------------------


async def enviar_consultoriafase2(session: AsyncSession, payload: dict[str, Any]) -> dict[str, Any]:
    """Send lead for consultoria phase 2 analysis via external webhook.

    Port of: app/api/admin/leads-chatwit/enviar-consultoriafase2/route.ts
    """
    lead_id = payload.get("leadID")
    if not lead_id:
        raise AnaliseServiceError("leadID não fornecido")

    webhook_url = settings.socialwise_webhook_url
    if not webhook_url:
        raise AnaliseServiceError("URL do webhook não configurada")

    # Fetch lead with relationships
    stmt = (
        select(LeadOabData)
        .options(
            selectinload(LeadOabData.lead),
            selectinload(LeadOabData.usuario_chatwit),
        )
        .where(LeadOabData.id == lead_id)
    )
    result = await session.execute(stmt)
    lead_oab = result.scalar_one_or_none()

    if not lead_oab:
        raise AnaliseServiceError("Lead não encontrado")

    # Fetch arquivos
    arquivos_stmt = select(ArquivoLeadOab).where(ArquivoLeadOab.lead_oab_data_id == lead_id)
    arquivos_result = await session.execute(arquivos_stmt)
    arquivos = list(arquivos_result.scalars().all())

    # Mark as waiting
    await session.execute(
        update(LeadOabData)
        .where(LeadOabData.id == lead_id)
        .values(aguardando_analise=True)
    )
    await session.commit()

    # Resolve espelho from biblioteca if linked
    espelho_biblioteca = None
    if lead_oab.espelho_biblioteca_id:
        bib_stmt = select(EspelhoBiblioteca).where(EspelhoBiblioteca.id == lead_oab.espelho_biblioteca_id)
        bib_result = await session.execute(bib_stmt)
        espelho_biblioteca = bib_result.scalar_one_or_none()
        if espelho_biblioteca:
            logger.info("espelho_biblioteca_found", nome=espelho_biblioteca.nome)

    # Format manuscript text
    texto_manuscrito = _format_text_value(lead_oab.prova_manuscrita, "Texto da Prova:")

    # Format espelho text — prioritize biblioteca
    texto_espelho = ""
    imagens_espelho: list[str] = []

    if espelho_biblioteca:
        texto_espelho = _format_text_value(
            espelho_biblioteca.texto_do_espelho, "Espelho da Prova (Biblioteca):"
        )
        imagens_espelho = _parse_espelho_images(espelho_biblioteca.espelho_correcao)
    elif lead_oab.texto_do_espelho:
        texto_espelho = _format_text_value(lead_oab.texto_do_espelho, "Espelho da Prova:")
        imagens_espelho = _parse_espelho_images(lead_oab.espelho_correcao)

    lead_name = lead_oab.nome_real or (lead_oab.lead.name if lead_oab.lead else None) or "Lead sem nome"
    lead_phone = lead_oab.lead.phone if lead_oab.lead else None

    # Build webhook payload
    request_payload: dict[str, Any] = {
        **payload,
        "analisesimulado": True,
        "leadID": lead_id,
        "nome": lead_name,
        "telefone": lead_phone,
        "textoManuscrito": texto_manuscrito,
        "textoEspelho": texto_espelho,
        "arquivos": [
            {"id": a.id, "url": a.data_url, "tipo": a.file_type, "nome": a.file_type}
            for a in arquivos
        ],
        "arquivos_pdf": (
            [{"id": lead_id, "url": lead_oab.pdf_unificado, "nome": "PDF Unificado"}]
            if lead_oab.pdf_unificado
            else []
        ),
        "arquivos_imagens_espelho": [
            {"id": f"{lead_id}-espelho-{i}", "url": url, "nome": f"Espelho {i + 1}"}
            for i, url in enumerate(imagens_espelho)
        ],
        "metadata": {
            "leadUrl": lead_oab.lead_url,
            "sourceId": lead_oab.lead.source_identifier if lead_oab.lead else None,
            "concluido": lead_oab.concluido,
            "fezRecurso": lead_oab.fez_recurso,
            "manuscritoProcessado": lead_oab.manuscrito_processado,
            "temEspelho": bool(lead_oab.espelho_correcao) or bool(espelho_biblioteca),
            "espelhoBibliotecaId": lead_oab.espelho_biblioteca_id,
            "espelhoBibliotecaNome": espelho_biblioteca.nome if espelho_biblioteca else None,
        },
    }

    logger.info(
        "consultoria_fase2_sending",
        lead_id=lead_id,
        tipo_espelho="biblioteca" if espelho_biblioteca else "individual",
        n_imagens=len(imagens_espelho),
    )

    # Fire-and-forget webhook
    async def _send_webhook() -> None:
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                resp = await client.post(webhook_url, json=request_payload)
                if resp.status_code >= 400:
                    logger.error("consultoria_webhook_error", status=resp.status_code)
                else:
                    logger.info("consultoria_webhook_sent", lead_id=lead_id)
        except Exception as exc:
            logger.error("consultoria_webhook_failed", error=str(exc))

    asyncio.create_task(_send_webhook())

    return {"success": True, "message": "Solicitação de consultoria fase 2 enviada com sucesso"}
