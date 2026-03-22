"""Service for generating contract documents (PDF, DOCX) and sending to clients."""

import io
import logging
from datetime import date
from decimal import Decimal
from typing import Optional
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession
from weasyprint import HTML

from domains.jusmonitoria.db.models.contrato import Contrato
from domains.jusmonitoria.db.repositories.contrato import ContratoRepository

logger = logging.getLogger(__name__)


def _format_currency(value: Optional[Decimal]) -> str:
    if value is None:
        return "-"
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _format_date(d: Optional[date]) -> str:
    if d is None:
        return "-"
    return d.strftime("%d/%m/%Y")


TIPO_LABELS = {
    "prestacao_servicos": "Prestação de Serviços",
    "honorarios_exito": "Honorários de Êxito",
    "misto": "Misto",
    "consultoria": "Consultoria",
    "contencioso": "Contencioso",
}


def _build_contract_html(contrato: Contrato) -> str:
    """Build the full HTML document for a contract, wrapping conteudo_html with header/footer."""
    client_name = contrato.client.full_name if contrato.client else "Cliente"
    tipo_label = TIPO_LABELS.get(contrato.tipo.value, contrato.tipo.value)

    body_content = contrato.conteudo_html or ""

    # If no rich text content, build from clausulas
    if not body_content.strip():
        parts = []
        if contrato.descricao:
            parts.append(f"<p>{contrato.descricao}</p>")
        if contrato.clausulas:
            for i, clausula in enumerate(contrato.clausulas, 1):
                titulo = clausula.get("titulo", f"Cláusula {i}")
                descricao = clausula.get("descricao", "")
                parts.append(
                    f'<h3>CLÁUSULA {i}ª — {titulo.upper()}</h3>'
                    f"<p>{descricao}</p>"
                )
        body_content = "\n".join(parts)

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<style>
  @page {{
    size: A4;
    margin: 2.5cm 2cm;
  }}
  body {{
    font-family: 'Georgia', 'Times New Roman', serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #1a1a1a;
  }}
  .header {{
    text-align: center;
    border-bottom: 2px solid #D4AF37;
    padding-bottom: 20px;
    margin-bottom: 30px;
  }}
  .header h1 {{
    font-size: 18pt;
    color: #0B0F19;
    margin: 0 0 5px 0;
    letter-spacing: 1px;
  }}
  .header .contract-number {{
    font-size: 10pt;
    color: #666;
    margin: 5px 0;
  }}
  .header .contract-type {{
    font-size: 11pt;
    color: #D4AF37;
    font-weight: bold;
  }}
  .parties {{
    background: #f8f7f4;
    border: 1px solid #e8e6e0;
    border-radius: 8px;
    padding: 20px;
    margin-bottom: 25px;
  }}
  .parties h2 {{
    font-size: 13pt;
    color: #0B0F19;
    margin: 0 0 12px 0;
    border-bottom: 1px solid #D4AF37;
    padding-bottom: 8px;
  }}
  .parties p {{
    margin: 5px 0;
    font-size: 11pt;
  }}
  .parties .label {{
    font-weight: bold;
    color: #555;
  }}
  .content {{
    margin-top: 20px;
  }}
  .content h2 {{
    font-size: 14pt;
    color: #0B0F19;
    margin-top: 25px;
  }}
  .content h3 {{
    font-size: 12pt;
    color: #333;
    margin-top: 20px;
    margin-bottom: 8px;
  }}
  .content p {{
    text-align: justify;
    margin: 8px 0;
  }}
  .content ul, .content ol {{
    margin: 8px 0;
    padding-left: 25px;
  }}
  .financial {{
    background: #f8f7f4;
    border: 1px solid #e8e6e0;
    border-radius: 8px;
    padding: 20px;
    margin: 25px 0;
  }}
  .financial h2 {{
    margin: 0 0 12px 0;
    font-size: 13pt;
    border-bottom: 1px solid #D4AF37;
    padding-bottom: 8px;
  }}
  .financial table {{
    width: 100%;
    border-collapse: collapse;
  }}
  .financial td {{
    padding: 6px 0;
    font-size: 11pt;
  }}
  .financial td:first-child {{
    font-weight: bold;
    color: #555;
    width: 40%;
  }}
  .signature {{
    margin-top: 60px;
    page-break-inside: avoid;
  }}
  .signature .date {{
    text-align: right;
    margin-bottom: 50px;
    font-size: 11pt;
  }}
  .signature .lines {{
    display: flex;
    justify-content: space-between;
    gap: 60px;
  }}
  .signature .line {{
    flex: 1;
    text-align: center;
  }}
  .signature .line hr {{
    border: none;
    border-top: 1px solid #333;
    margin-bottom: 8px;
  }}
  .signature .line p {{
    font-size: 10pt;
    margin: 3px 0;
    color: #333;
  }}
  .footer {{
    margin-top: 40px;
    padding-top: 15px;
    border-top: 1px solid #e8e6e0;
    font-size: 9pt;
    color: #999;
    text-align: center;
  }}
</style>
</head>
<body>

<div class="header">
  <h1>{contrato.titulo}</h1>
  <p class="contract-number">{contrato.numero_contrato}</p>
  <p class="contract-type">{tipo_label}</p>
</div>

<div class="parties">
  <h2>PARTES DO CONTRATO</h2>
  <p><span class="label">CONTRATANTE:</span> {client_name}</p>
  <p><span class="label">CONTRATADA:</span> Escritório de Advocacia</p>
  {f'<p><span class="label">ADVOGADO RESPONSÁVEL:</span> {contrato.assigned_user.full_name}</p>' if contrato.assigned_user else ''}
</div>

<div class="financial">
  <h2>CONDIÇÕES FINANCEIRAS</h2>
  <table>
    {f'<tr><td>Valor Mensal:</td><td>{_format_currency(contrato.valor_mensal)}</td></tr>' if contrato.valor_mensal else ''}
    {f'<tr><td>Valor Total:</td><td>{_format_currency(contrato.valor_total)}</td></tr>' if contrato.valor_total else ''}
    {f'<tr><td>Entrada:</td><td>{_format_currency(contrato.valor_entrada)}</td></tr>' if contrato.valor_entrada else ''}
    {f'<tr><td>Percentual de Êxito:</td><td>{contrato.percentual_exito}%</td></tr>' if contrato.percentual_exito else ''}
    {f'<tr><td>Data de Início:</td><td>{_format_date(contrato.data_inicio)}</td></tr>' if contrato.data_inicio else ''}
    {f'<tr><td>Data de Vencimento:</td><td>{_format_date(contrato.data_vencimento)}</td></tr>' if contrato.data_vencimento else ''}
  </table>
</div>

<div class="content">
{body_content}
</div>

<div class="signature">
  <p class="date">Data: ____/____/________</p>
  <div class="lines" style="display: flex; justify-content: space-between;">
    <div class="line" style="flex: 1; text-align: center;">
      <hr>
      <p><strong>{client_name}</strong></p>
      <p>Contratante</p>
    </div>
    <div style="width: 60px;"></div>
    <div class="line" style="flex: 1; text-align: center;">
      <hr>
      <p><strong>{contrato.assigned_user.full_name if contrato.assigned_user else 'Advogado Responsável'}</strong></p>
      <p>Contratada</p>
    </div>
  </div>
</div>

<div class="footer">
  Documento gerado automaticamente pelo JusMonitorIA
</div>

</body>
</html>"""
    return html


class ContratoDocumentService:
    """Service for generating PDF/DOCX contract documents."""

    def __init__(self, session: AsyncSession, tenant_id: UUID):
        self.session = session
        self.tenant_id = tenant_id
        self.contrato_repo = ContratoRepository(session, tenant_id)

    async def _get_contrato(self, contrato_id: UUID) -> Contrato:
        contrato = await self.contrato_repo.get(contrato_id)
        if not contrato:
            raise ValueError("Contrato não encontrado")
        return contrato

    async def generate_pdf(self, contrato_id: UUID) -> bytes:
        """Generate a PDF document from the contract."""
        contrato = await self._get_contrato(contrato_id)
        html_content = _build_contract_html(contrato)

        pdf_bytes = HTML(string=html_content).write_pdf()

        logger.info(
            "contrato_pdf_gerado",
            extra={
                "contrato_id": str(contrato_id),
                "tenant_id": str(self.tenant_id),
                "size_bytes": len(pdf_bytes),
            },
        )
        return pdf_bytes

    async def generate_docx(self, contrato_id: UUID) -> bytes:
        """Generate a DOCX document from the contract."""
        contrato = await self._get_contrato(contrato_id)
        client_name = contrato.client.full_name if contrato.client else "Cliente"
        tipo_label = TIPO_LABELS.get(contrato.tipo.value, contrato.tipo.value)

        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Georgia"
        font.size = Pt(12)
        font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style.paragraph_format.line_spacing = 1.5

        # Title
        title = doc.add_heading(contrato.titulo, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x0B, 0x0F, 0x19)

        # Contract number and type
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{contrato.numero_contrato}\n{tipo_label}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # spacer

        # Parties
        doc.add_heading("PARTES DO CONTRATO", level=2)
        doc.add_paragraph(f"CONTRATANTE: {client_name}")
        doc.add_paragraph("CONTRATADA: Escritório de Advocacia")
        if contrato.assigned_user:
            doc.add_paragraph(f"ADVOGADO RESPONSÁVEL: {contrato.assigned_user.full_name}")

        # Financial
        doc.add_heading("CONDIÇÕES FINANCEIRAS", level=2)
        if contrato.valor_mensal:
            doc.add_paragraph(f"Valor Mensal: {_format_currency(contrato.valor_mensal)}")
        if contrato.valor_total:
            doc.add_paragraph(f"Valor Total: {_format_currency(contrato.valor_total)}")
        if contrato.valor_entrada:
            doc.add_paragraph(f"Entrada: {_format_currency(contrato.valor_entrada)}")
        if contrato.percentual_exito:
            doc.add_paragraph(f"Percentual de Êxito: {contrato.percentual_exito}%")
        if contrato.data_inicio:
            doc.add_paragraph(f"Data de Início: {_format_date(contrato.data_inicio)}")
        if contrato.data_vencimento:
            doc.add_paragraph(f"Data de Vencimento: {_format_date(contrato.data_vencimento)}")

        # Content
        if contrato.conteudo_html and contrato.conteudo_html.strip():
            doc.add_heading("TERMOS E CONDIÇÕES", level=2)
            # Strip HTML tags for DOCX (simple conversion)
            import re

            text = re.sub(r"<br\s*/?>", "\n", contrato.conteudo_html)
            text = re.sub(r"</p>", "\n\n", text)
            text = re.sub(r"</h[1-6]>", "\n\n", text)
            text = re.sub(r"<h[1-6][^>]*>", "", text)
            text = re.sub(r"<li[^>]*>", "  • ", text)
            text = re.sub(r"</li>", "\n", text)
            text = re.sub(r"<[^>]+>", "", text)
            text = text.strip()

            for paragraph_text in text.split("\n\n"):
                paragraph_text = paragraph_text.strip()
                if paragraph_text:
                    p = doc.add_paragraph(paragraph_text)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif contrato.clausulas:
            doc.add_heading("TERMOS E CONDIÇÕES", level=2)
            for i, clausula in enumerate(contrato.clausulas, 1):
                titulo = clausula.get("titulo", f"Cláusula {i}")
                descricao = clausula.get("descricao", "")

                h = doc.add_heading(f"CLÁUSULA {i}ª — {titulo.upper()}", level=3)
                for run in h.runs:
                    run.font.size = Pt(12)

                if descricao:
                    p = doc.add_paragraph(descricao)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph("Data: ____/____/________")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph("_" * 40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(client_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        p = doc.add_paragraph("Contratante")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph("_" * 40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lawyer_name = contrato.assigned_user.full_name if contrato.assigned_user else "Advogado Responsável"
        p = doc.add_paragraph(lawyer_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        p = doc.add_paragraph("Contratada")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph("Documento gerado automaticamente pelo JusMonitorIA")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Write to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()

        logger.info(
            "contrato_docx_gerado",
            extra={
                "contrato_id": str(contrato_id),
                "tenant_id": str(self.tenant_id),
                "size_bytes": len(docx_bytes),
            },
        )
        return docx_bytes

    async def send_to_client_chatwit(
        self,
        contrato_id: UUID,
        pdf_url: str,
    ) -> dict:
        """Send contract PDF link to client via ChatWoot (WhatsApp)."""
        contrato = await self._get_contrato(contrato_id)

        if not contrato.client:
            raise ValueError("Cliente não encontrado")

        chatwit_contact_id = contrato.client.chatwit_contact_id
        if not chatwit_contact_id:
            raise ValueError("Cliente não possui ID de contato no ChatWoot/WhatsApp")

        from domains.jusmonitoria.services.chatwit_client import get_chatwit_client

        client = get_chatwit_client()

        message = (
            f"Olá {contrato.client.full_name},\n\n"
            f"Segue o contrato *{contrato.titulo}* ({contrato.numero_contrato}) "
            f"para sua análise e assinatura.\n\n"
            f"📄 Acesse o documento: {pdf_url}\n\n"
            f"Após revisar, você pode:\n"
            f"• Assinar digitalmente via gov.br\n"
            f"• Imprimir, assinar e nos enviar uma cópia\n\n"
            f"Qualquer dúvida, estamos à disposição."
        )

        result = await client.send_message(
            contact_id=chatwit_contact_id,
            message=message,
            channel="whatsapp",
        )

        logger.info(
            "contrato_enviado_chatwit",
            extra={
                "contrato_id": str(contrato_id),
                "contact_id": chatwit_contact_id,
                "tenant_id": str(self.tenant_id),
            },
        )
        return result

    async def send_to_client_email(
        self,
        contrato_id: UUID,
        pdf_url: str,
    ) -> bool:
        """Send contract PDF link to client via email."""
        contrato = await self._get_contrato(contrato_id)

        if not contrato.client:
            raise ValueError("Cliente não encontrado")

        client_email = contrato.client.email
        if not client_email:
            raise ValueError("Cliente não possui e-mail cadastrado")

        from domains.jusmonitoria.services.email_service import EmailService

        client_name = contrato.client.full_name

        subject = f"Contrato {contrato.numero_contrato} — {contrato.titulo}"

        html_content = f"""
        <div style="font-family: Georgia, serif; max-width: 600px; margin: 0 auto; color: #1a1a1a;">
            <div style="text-align: center; border-bottom: 2px solid #D4AF37; padding-bottom: 20px; margin-bottom: 25px;">
                <h2 style="color: #0B0F19; margin: 0;">Contrato para Assinatura</h2>
            </div>

            <p>Prezado(a) <strong>{client_name}</strong>,</p>

            <p>Encaminhamos o contrato <strong>{contrato.titulo}</strong>
            (Nº {contrato.numero_contrato}) para sua análise e assinatura.</p>

            <div style="text-align: center; margin: 30px 0;">
                <a href="{pdf_url}"
                   style="background-color: #D4AF37; color: #0B0F19; padding: 14px 28px;
                          text-decoration: none; border-radius: 6px; font-weight: bold;
                          font-size: 16px; display: inline-block;">
                    Baixar Contrato (PDF)
                </a>
            </div>

            <p>Após revisar o documento, você pode:</p>
            <ul style="line-height: 2;">
                <li>Assinar digitalmente via <strong>gov.br</strong></li>
                <li>Imprimir, assinar e nos enviar uma cópia digitalizada</li>
            </ul>

            <hr style="border: 1px solid #e8e6e0; margin: 30px 0;">

            <p style="font-size: 12px; color: #999;">
                Este é um e-mail automático enviado pelo JusMonitorIA.
                Em caso de dúvidas, entre em contato com seu advogado responsável.
            </p>
        </div>
        """

        text_content = (
            f"Prezado(a) {client_name},\n\n"
            f"Encaminhamos o contrato {contrato.titulo} "
            f"(Nº {contrato.numero_contrato}) para sua análise e assinatura.\n\n"
            f"Acesse o documento: {pdf_url}\n\n"
            f"Após revisar, você pode:\n"
            f"- Assinar digitalmente via gov.br\n"
            f"- Imprimir, assinar e nos enviar uma cópia\n\n"
            f"Qualquer dúvida, estamos à disposição."
        )

        result = await EmailService.send_email(
            to_email=client_email,
            subject=subject,
            html_content=html_content,
            text_content=text_content,
        )

        logger.info(
            "contrato_enviado_email",
            extra={
                "contrato_id": str(contrato_id),
                "email": client_email,
                "tenant_id": str(self.tenant_id),
                "success": result,
            },
        )
        return result
